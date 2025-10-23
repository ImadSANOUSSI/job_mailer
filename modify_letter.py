import os
import argparse
from docx import Document
from utils import get_env, get_openai_client, logger, safe_filename


def read_template(path: str) -> str:
    d = Document(path)
    return "\n".join(p.text for p in d.paragraphs)


def save_docx_from_text(text: str, out_path: str) -> None:
    d = Document()
    for line in text.splitlines():
        d.add_paragraph(line)
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    d.save(out_path)


def ask_model(template_text: str, job_post_text: str, extra_instructions: str = "") -> str:
    client = get_openai_client()
    system = (
        "Tu es un assistant expert en rédaction de lettres de motivation en français. "
        "Adapte le modèle précisément à l'offre d'emploi, ton professionnel et concis (~1 page)."
    )
    user = (
        f"Modèle de lettre:\n{template_text}\n\n"
        f"Offre d'emploi:\n{job_post_text}\n\n"
        "Consignes: génère la lettre finale en français. Conserve les compétences clés, "
        "ajoute une motivation claire et adapte au poste."
    )
    if extra_instructions:
        user += "\n\n" + extra_instructions
    r = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
        temperature=0.2,
        max_tokens=1200,
    )
    return (r.choices[0].message.content or "").strip()


def generate_for_post(job_post_text: str, out_folder: str, filename_prefix: str = "lettre", extra_instructions: str = "") -> str:
    template_path = get_env("TEMPLATE_PATH", required=True)
    if not os.path.exists(template_path):
        raise FileNotFoundError("Template file not found")
    template_text = read_template(template_path)
    letter_text = ask_model(template_text, job_post_text, extra_instructions)
    safe = safe_filename(filename_prefix)
    out_path = os.path.join(out_folder, f"{safe}.docx")
    save_docx_from_text(letter_text, out_path)
    return out_path


def main() -> None:
    p = argparse.ArgumentParser(description="Generate a personalized French cover letter from a job post.")
    p.add_argument("--jobfile", required=True)
    p.add_argument("--out", required=True)
    p.add_argument("--prefix", default="lettre")
    p.add_argument("--extra", default="")
    a = p.parse_args()
    with open(a.jobfile, encoding="utf-8") as f:
        job_text = f.read()
    path = generate_for_post(job_text, a.out, a.prefix, a.extra)
    logger.info(f"Saved: {path}")


if __name__ == "__main__":
    main()
